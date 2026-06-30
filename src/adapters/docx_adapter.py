"""
DOCX Adapter — extracts candidate data from resume DOCX files.

Uses python-docx for text extraction.
Handles: corrupted file, empty doc, tables vs paragraphs.
"""

import os
import re
from src.adapters.base_adapter import BaseAdapter
from src.schema.source_record import SourceRecord
from src.logger import logger


class DocxAdapter(BaseAdapter):
    """Adapter for resume DOCX files."""

    def __init__(self):
        super().__init__()
        self.source_name = "resume_docx"
        self.supported_extensions = [".docx"]

    def extract(self, input_data):
        """Extract candidate data from a DOCX resume."""
        if not input_data or not os.path.exists(input_data):
            logger.warning("ingest", self.source_name, f"File not found: {input_data}")
            return [self._empty_record_with_error(f"File not found: {input_data}")]

        if os.path.getsize(input_data) == 0:
            logger.warning("ingest", self.source_name, "DOCX file is empty")
            return [self._empty_record_with_error("DOCX file is empty")]

        try:
            from docx import Document
            doc = Document(input_data)

            # Extract text from paragraphs
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            paragraphs.append(cell_text)

            text = "\n".join(paragraphs)
        except Exception as e:
            logger.error("ingest", self.source_name, f"Failed to read DOCX: {str(e)}")
            return [self._empty_record_with_error(f"Failed to read DOCX: {str(e)}")]

        if not text.strip():
            logger.warning("ingest", self.source_name, "DOCX contains no text")
            return [self._empty_record_with_error("DOCX contains no text")]

        record = self._parse_text(text)
        record.raw_text = text
        logger.info("ingest", self.source_name, f"Extracted text from DOCX ({len(text)} chars)")
        return [record]

    def _parse_text(self, text):
        """Parse raw text from a DOCX resume. Same heuristics as PDF adapter."""
        record = SourceRecord(
            source_name=self.source_name,
            extraction_method="regex",
        )

        # Extract emails
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        record.emails = list(set(re.findall(email_pattern, text)))

        # Extract phone numbers
        phone_pattern = r'[\+]?[\d\s\-\(\)]{7,15}'
        raw_phones = re.findall(phone_pattern, text)
        record.phones = [p.strip() for p in raw_phones if len(re.sub(r'\D', '', p)) >= 7]

        # Name — first non-empty line
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        if lines:
            first_line = lines[0]
            if len(first_line) < 60 and not re.search(r'(resume|curriculum|vitae|cv)', first_line, re.IGNORECASE):
                record.full_name = first_line

        # Links
        url_pattern = r'https?://[^\s<>\"\')\]]*'
        urls = re.findall(url_pattern, text)
        for url in urls:
            lower = url.lower()
            if "github.com" in lower:
                record.links["github"] = url
            elif "linkedin.com" in lower:
                record.links["linkedin"] = url

        # Skills section
        skills_text = self._extract_section(text, ["skills", "technical skills", "technologies"])
        if skills_text:
            skills = re.split(r'[,|•·\n]', skills_text)
            record.skills = [s.strip() for s in skills if s.strip() and len(s.strip()) < 50]

        # Location
        loc_match = re.search(r'([A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s*[A-Z]{2})', text)
        if loc_match:
            record.location = loc_match.group(1).strip()

        return record

    def _extract_section(self, text, section_names):
        """Extract text under a section header."""
        lines = text.split("\n")
        capturing = False
        captured = []
        stop_headers = [
            "experience", "work experience", "education",
            "projects", "certifications", "summary",
            "objective", "references", "awards",
        ]

        for line in lines:
            stripped = line.strip().lower()
            if not capturing:
                for name in section_names:
                    if stripped.startswith(name) or stripped == name:
                        capturing = True
                        break
            else:
                is_new = any(stripped.startswith(h) or stripped == h for h in stop_headers)
                if is_new:
                    break
                if line.strip():
                    captured.append(line.strip())

        return " ".join(captured)
